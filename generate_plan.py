from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Pt

doc = Document()

title = doc.add_heading('Peptide Store MVP Architecture & Technical Plan', level=1)
title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

p = doc.add_paragraph(
    "This document outlines a lightweight MVP architecture for a peptide store "
    "with WhatsApp ordering, real-time inventory management, and a simple admin panel."
)

sections = [
    ("Core Goals", [
        "Simple static storefront",
        "Inventory editable by non-technical staff",
        "WhatsApp-based ordering",
        "Low hosting and maintenance cost",
        "Mobile-first design",
        "Easy scaling later"
    ]),
    ("Recommended Stack", [
        "Frontend: React + Vite + TailwindCSS",
        "Hosting: Vercel",
        "Database: Supabase",
        "Authentication: Supabase Auth",
        "Admin Panel: Protected React routes",
        "Storage: Supabase Storage"
    ]),
    ("Main User Flow", [
        "Customer visits website",
        "Customer browses products",
        "Customer clicks WhatsApp order button",
        "Pre-filled WhatsApp message opens",
        "Store staff handles order manually"
    ]),
    ("Admin Flow", [
        "Employee logs into admin panel",
        "Employee edits inventory quantities",
        "Website updates automatically",
        "Employee can hide/show products",
        "Employee can update pricing"
    ]),
]

for heading, items in sections:
    doc.add_heading(heading, level=2)
    for item in items:
        doc.add_paragraph(item, style='List Bullet')

doc.add_heading("Suggested Database Schema", level=2)

schema = """
Table: products

id UUID PRIMARY KEY
name TEXT
description TEXT
price DECIMAL
stock INTEGER
category TEXT
image_url TEXT
visible BOOLEAN DEFAULT true
featured BOOLEAN DEFAULT false
created_at TIMESTAMP
updated_at TIMESTAMP
"""

doc.add_paragraph(schema)

doc.add_heading("Recommended Folder Structure", level=2)

folder_structure = """
/src
  /components
    Navbar.jsx
    ProductCard.jsx
    ProductGrid.jsx
    WhatsAppButton.jsx
    AdminSidebar.jsx

  /pages
    Home.jsx
    Products.jsx
    Admin.jsx
    Login.jsx

  /services
    supabase.js
    products.js

  /hooks
    useProducts.js

  /styles
    globals.css

  App.jsx
  main.jsx
"""

doc.add_paragraph(folder_structure)

doc.add_heading("Supabase Client Example", level=2)

code1 = """
import { createClient } from '@supabase/supabase-js'

const supabaseUrl = import.meta.env.VITE_SUPABASE_URL
const supabaseKey = import.meta.env.VITE_SUPABASE_ANON_KEY

export const supabase = createClient(
  supabaseUrl,
  supabaseKey
)
"""

doc.add_paragraph(code1)

doc.add_heading("Fetch Products Example", level=2)

code2 = """
const { data, error } = await supabase
  .from('products')
  .select('*')
  .eq('visible', true)
"""

doc.add_paragraph(code2)

doc.add_heading("WhatsApp Order Integration", level=2)

code3 = """
const message = `Hello, I'm interested in:
Product: ${product.name}
Price: $${product.price}`

const whatsappUrl =
  `https://wa.me/18095551234?text=${encodeURIComponent(message)}`
"""

doc.add_paragraph(code3)

doc.add_heading("Admin Panel Features", level=2)

admin_features = [
    "Login authentication",
    "Update inventory",
    "Toggle product visibility",
    "Edit prices",
    "Upload product images",
    "Search/filter products",
    "Mobile responsive layout"
]

for item in admin_features:
    doc.add_paragraph(item, style='List Bullet')

doc.add_heading("Recommended UI Style", level=2)

ui_items = [
    "Dark modern aesthetic",
    "Laboratory-inspired visuals",
    "Minimal clean cards",
    "Large mobile-friendly buttons",
    "Sticky WhatsApp contact button"
]

for item in ui_items:
    doc.add_paragraph(item, style='List Bullet')

doc.add_heading("Deployment Plan", level=2)

deployment = [
    "Push frontend to GitHub",
    "Connect GitHub repository to Vercel",
    "Add environment variables",
    "Deploy automatically on every push",
    "Connect Supabase project"
]

for item in deployment:
    doc.add_paragraph(item, style='List Bullet')

doc.add_heading("Estimated Timeline", level=2)

timeline = [
    "Day 1: Frontend setup + Supabase configuration",
    "Day 2: Product pages + inventory sync",
    "Day 3: Admin panel + authentication",
    "Day 4: Mobile optimization + deployment"
]

for item in timeline:
    doc.add_paragraph(item, style='List Bullet')

doc.add_heading("Future Features", level=2)

future = [
    "Customer accounts",
    "Online payment processing",
    "Order tracking",
    "Analytics dashboard",
    "Inventory alerts",
    "Multi-language support"
]

for item in future:
    doc.add_paragraph(item, style='List Bullet')

path = "/mnt/data/peptide_store_mvp_plan.docx"
doc.save(path)

print(f"Document saved to: {path}")
